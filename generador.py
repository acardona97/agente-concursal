"""
generador.py v2 — Generador de documentos con estilo Quarta
Parsea créditos directamente de la solicitud del usuario con regex,
sin depender de un segundo llamado a la API.
"""

import os
import re
import json
from pathlib import Path
from datetime import datetime
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_NAVY   = "1B2A4A"
COLOR_GOLD   = "C9A84C"
COLOR_LIGHT  = "EEF1F6"
COLOR_WHITE  = "FFFFFF"
COLOR_TEXT   = "1A1A1A"
COLOR_BORDER = "C5CCDA"

LOGO_PATH = "/mnt/skills/user/quarta-propuestas/quarta_logo_oficial.png"

TRIGGERS_EXCEL = [
    "calificaci", "graduaci", "tabla de cr", "lista de acreedor",
    "proyecto de calificaci", "excel", "xlsx", "tabla de votos"
]
TRIGGERS_WORD = [
    "memorial", "escrito", "recurso", "objecion", "objeción",
    "solicitud", "peticion", "petición", "demanda", "apertura", "acuerdo"
]


def detectar_tipo_documento(solicitud: str) -> str:
    s = solicitud.lower()
    for t in TRIGGERS_EXCEL:
        if t in s:
            return "excel"
    for t in TRIGGERS_WORD:
        if t in s:
            return "word"
    return "texto"


def _limpiar_valor(valor_str: str) -> int:
    """Convierte '$150.000.000' o '150000000' a int."""
    v = re.sub(r'[^\d]', '', valor_str)
    try:
        return int(v)
    except Exception:
        return 0


def _normalizar_clase(clase_raw: str) -> str:
    clases = ["Primera", "Segunda", "Tercera", "Cuarta", "Quinta", "Quirografario", "Subordinado"]
    c = clase_raw.strip().lower()
    for cl in clases:
        if cl.lower() in c:
            return cl
    return clase_raw.strip().title()


def _parsear_creditos(texto: str, nombre_proceso: str) -> dict:
    """
    Parsea créditos desde texto usando regex.
    Busca líneas con formato:
    N. Acreedor - Clase - Concepto - $Valor - Garantia
    """
    creditos = []

    # Patron principal: número. texto - texto - texto - $valor - texto
    patron = re.compile(
        r'(\d+)[.)]\s*(.+?)\s*[-–]\s*(.+?)\s*[-–]\s*(.+?)\s*[-–]\s*\$?([\d.,]+)\s*[-–]\s*(.+?)(?:\n|$)',
        re.MULTILINE
    )
    matches = patron.findall(texto)

    for m in matches:
        num, acreedor, clase, concepto, valor_str, garantia = m
        creditos.append({
            "numero": int(num),
            "acreedor": acreedor.strip(),
            "clase": _normalizar_clase(clase),
            "concepto": concepto.strip(),
            "valor": _limpiar_valor(valor_str),
            "garantia": garantia.strip(),
            "observaciones": ""
        })

    # Buscar nombre del deudor
    deudor = nombre_proceso
    patrones_deudor = [
        r'proceso de (?:reorganizaci[oó]n de |liquidaci[oó]n de )?([A-ZÁÉÍÓÚÑ][^,\n]+?)(?:,|\n|con los)',
        r'deudor[:\s]+([A-ZÁÉÍÓÚÑ][^\n,]+)',
        r'(?:empresa|sociedad)[:\s]+([A-ZÁÉÍÓÚÑ][^\n,]+)',
    ]
    for p in patrones_deudor:
        m = re.search(p, texto, re.IGNORECASE)
        if m:
            deudor = m.group(1).strip()
            break

    if not creditos:
        creditos = [{
            "numero": 1, "acreedor": "[COMPLETAR]", "clase": "Primera",
            "concepto": "Capital", "valor": 0, "garantia": "Sin garantía",
            "observaciones": "Completar manualmente"
        }]

    return {
        "deudor": deudor,
        "fecha": datetime.now().strftime("%d/%m/%Y"),
        "creditos": creditos
    }


def generar_excel_calificacion(solicitud: str, respuesta_agente: str,
                                nombre_proceso: str = "Proceso") -> str:
    # Parsear desde la solicitud original (tiene los datos más limpios)
    data = _parsear_creditos(solicitud, nombre_proceso)

    # Si no encontró créditos en la solicitud, intentar con la respuesta
    if len(data["creditos"]) == 1 and data["creditos"][0]["acreedor"] == "[COMPLETAR]":
        data = _parsear_creditos(respuesta_agente, nombre_proceso)

    deudor   = data["deudor"]
    fecha    = data["fecha"]
    creditos = data["creditos"]

    # Ordenar por clase
    orden = {"primera": 1, "segunda": 2, "tercera": 3, "cuarta": 4,
             "quinta": 5, "quirografario": 6, "subordinado": 7}
    creditos = sorted(creditos, key=lambda x: orden.get(x["clase"].lower(), 99))

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Calificación y Graduación"

    # Estilos
    FNAV = Font(name="Calibri", size=10, bold=True, color=COLOR_WHITE)
    FBOD = Font(name="Calibri", size=10, color=COLOR_TEXT)
    FTIT = Font(name="Calibri", size=13, bold=True, color=COLOR_NAVY)
    FTOT = Font(name="Calibri", size=10, bold=True, color=COLOR_NAVY)
    FNAV_FILL = PatternFill("solid", fgColor=COLOR_NAVY)
    FGOLD     = PatternFill("solid", fgColor=COLOR_GOLD)
    FLIGHT    = PatternFill("solid", fgColor=COLOR_LIGHT)
    FWHITE    = PatternFill("solid", fgColor=COLOR_WHITE)
    FTOT_FILL = PatternFill("solid", fgColor="D6DCE9")
    BTN = Border(
        left=Side(style="thin", color=COLOR_BORDER),
        right=Side(style="thin", color=COLOR_BORDER),
        top=Side(style="thin", color=COLOR_BORDER),
        bottom=Side(style="thin", color=COLOR_BORDER)
    )
    AC = Alignment(horizontal="center", vertical="center", wrap_text=True)
    AL = Alignment(horizontal="left",   vertical="center", wrap_text=True)
    AR = Alignment(horizontal="right",  vertical="center", wrap_text=True)

    # Barra dorada
    ws.row_dimensions[1].height = 6
    for c in range(1, 10):
        ws.cell(1, c).fill = FGOLD

    # Título
    ws.merge_cells("A2:I3")
    t = ws["A2"]
    t.value = "PROYECTO DE CALIFICACIÓN Y GRADUACIÓN DE CRÉDITOS"
    t.font = FTIT; t.alignment = AC; t.fill = FWHITE
    ws.row_dimensions[2].height = 22; ws.row_dimensions[3].height = 14

    # Subtítulo
    ws.merge_cells("A4:I4")
    s = ws["A4"]
    s.value = f"Proceso: {deudor.upper()}   |   Fecha: {fecha}   |   Ley 1116 de 2006"
    s.font = Font(name="Calibri", size=9, italic=True, color="555555")
    s.alignment = AC; s.fill = FLIGHT
    ws.row_dimensions[4].height = 15
    ws.row_dimensions[5].height = 8

    # Encabezados
    headers = ["N°", "Acreedor", "Clase / Prelación", "Concepto",
               "Valor Reconocido ($)", "Garantía", "Voto (Sí/No)",
               "% Participación", "Observaciones"]
    widths  = [5, 28, 18, 16, 20, 16, 12, 16, 28]
    for i, (h, w) in enumerate(zip(headers, widths), 1):
        c = ws.cell(6, i)
        c.value = h; c.font = FNAV; c.fill = FNAV_FILL
        c.alignment = AC; c.border = BTN
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.row_dimensions[6].height = 36

    # Filas de datos
    primera_fila = 7
    n = len(creditos)
    for idx, cr in enumerate(creditos):
        row = primera_fila + idx
        fill = FLIGHT if idx % 2 == 0 else FWHITE
        ws.row_dimensions[row].height = 20

        total_row = primera_fila + n

        datos = [
            cr["numero"], cr["acreedor"], cr["clase"], cr["concepto"],
            cr["valor"], cr["garantia"], "",
            f"=E{row}/E{total_row}" if cr["valor"] else "",
            cr["observaciones"]
        ]
        for col, val in enumerate(datos, 1):
            cell = ws.cell(row, col)
            cell.value = val; cell.border = BTN; cell.fill = fill
            cell.font = FBOD
            if col in [1, 3, 6, 7]:
                cell.alignment = AC
            elif col == 8:
                cell.alignment = AC
                cell.number_format = "0.00%"
            else:
                cell.alignment = AL
            if col == 5:
                cell.number_format = "$#,##0"

    # Fila total
    total_row = primera_fila + n
    ws.row_dimensions[total_row].height = 22
    ws.merge_cells(f"A{total_row}:D{total_row}")
    tl = ws[f"A{total_row}"]
    tl.value = "TOTAL CRÉDITOS RECONOCIDOS"
    tl.font = FTOT; tl.fill = FTOT_FILL; tl.alignment = AR; tl.border = BTN

    tv = ws[f"E{total_row}"]
    tv.value = f"=SUM(E{primera_fila}:E{total_row-1})"
    tv.font = Font(name="Calibri", size=10, bold=True, color=COLOR_NAVY)
    tv.fill = FTOT_FILL; tv.number_format = "$#,##0"; tv.alignment = AR; tv.border = BTN
    for col in [6, 7, 8, 9]:
        ws.cell(total_row, col).fill = FTOT_FILL
        ws.cell(total_row, col).border = BTN

    # Hoja 2: Resumen por clases
    ws2 = wb.create_sheet("Resumen por Clases")
    ws2.column_dimensions["A"].width = 25
    ws2.column_dimensions["B"].width = 22
    ws2.column_dimensions["C"].width = 15

    for c in range(1, 4):
        ws2.cell(1, c).fill = FGOLD
    ws2.row_dimensions[1].height = 6

    ws2.merge_cells("A2:C3")
    ws2["A2"].value = "RESUMEN POR CLASE DE CRÉDITO"
    ws2["A2"].font = Font(name="Calibri", size=13, bold=True, color=COLOR_NAVY)
    ws2["A2"].alignment = AC
    ws2.row_dimensions[2].height = 20

    ws2.merge_cells("A4:C4")
    ws2["A4"].value = f"{deudor.upper()} — {fecha}"
    ws2["A4"].font = Font(name="Calibri", size=9, italic=True, color="666666")
    ws2["A4"].alignment = AC; ws2["A4"].fill = FLIGHT
    ws2.row_dimensions[4].height = 15; ws2.row_dimensions[5].height = 8

    h2 = ["Clase / Prelación", "Total Reconocido ($)", "% del Total"]
    for i, h in enumerate(h2, 1):
        c = ws2.cell(6, i)
        c.value = h; c.font = FNAV; c.fill = FNAV_FILL
        c.alignment = AC; c.border = BTN
    ws2.row_dimensions[6].height = 28

    from collections import defaultdict
    por_clase = defaultdict(float)
    for cr in creditos:
        por_clase[cr["clase"]] += float(cr["valor"])

    orden_display = ["Primera","Segunda","Tercera","Cuarta","Quinta","Quirografario","Subordinado"]
    clases_doc = [c for c in orden_display if c in por_clase]
    clases_doc += [c for c in por_clase if c not in orden_display]

    fila_r = 7
    total_r = fila_r + len(clases_doc)
    for i, clase in enumerate(clases_doc):
        fill = FLIGHT if i % 2 == 0 else FWHITE
        ws2.cell(fila_r, 1).value = clase
        ws2.cell(fila_r, 1).font = FBOD; ws2.cell(fila_r, 1).fill = fill
        ws2.cell(fila_r, 1).alignment = AL; ws2.cell(fila_r, 1).border = BTN

        ws2.cell(fila_r, 2).value = por_clase[clase]
        ws2.cell(fila_r, 2).font = FBOD; ws2.cell(fila_r, 2).fill = fill
        ws2.cell(fila_r, 2).number_format = "$#,##0"
        ws2.cell(fila_r, 2).alignment = AR; ws2.cell(fila_r, 2).border = BTN

        ws2.cell(fila_r, 3).value = f"=B{fila_r}/B{total_r}"
        ws2.cell(fila_r, 3).font = FBOD; ws2.cell(fila_r, 3).fill = fill
        ws2.cell(fila_r, 3).number_format = "0.00%"
        ws2.cell(fila_r, 3).alignment = AC; ws2.cell(fila_r, 3).border = BTN
        ws2.row_dimensions[fila_r].height = 18
        fila_r += 1

    ws2.cell(total_r, 1).value = "TOTAL"
    ws2.cell(total_r, 1).font = FTOT; ws2.cell(total_r, 1).fill = FTOT_FILL
    ws2.cell(total_r, 1).alignment = AR; ws2.cell(total_r, 1).border = BTN

    ws2.cell(total_r, 2).value = f"=SUM(B7:B{total_r-1})"
    ws2.cell(total_r, 2).font = Font(name="Calibri", size=10, bold=True, color=COLOR_NAVY)
    ws2.cell(total_r, 2).fill = FTOT_FILL; ws2.cell(total_r, 2).number_format = "$#,##0"
    ws2.cell(total_r, 2).alignment = AR; ws2.cell(total_r, 2).border = BTN

    ws2.cell(total_r, 3).value = "100.00%"
    ws2.cell(total_r, 3).font = FTOT; ws2.cell(total_r, 3).fill = FTOT_FILL
    ws2.cell(total_r, 3).number_format = "0.00%"
    ws2.cell(total_r, 3).alignment = AC; ws2.cell(total_r, 3).border = BTN
    ws2.row_dimensions[total_r].height = 22

    ws2.cell(total_r + 2, 1).value = "Documento generado por Agente Concursal · Quarta Acompañamiento Legal S.A.S."
    ws2.cell(total_r + 2, 1).font = Font(name="Calibri", size=8, italic=True, color="999999")
    ws2.merge_cells(f"A{total_r+2}:C{total_r+2}")

    # Guardar
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^\w\s-]", "", deudor).strip().replace(" ", "_")
    out = Path("./documentos_generados")
    out.mkdir(exist_ok=True)
    ruta = out / f"Calificacion_{safe}_{ts}.xlsx"
    wb.save(str(ruta))
    return str(ruta)


def generar_word_escrito(respuesta_agente: str, nombre_proceso: str = "Proceso") -> str:
    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin    = Cm(3.0)
    section.bottom_margin = Cm(2.75)
    section.left_margin   = Cm(2.25)
    section.right_margin  = Cm(2.34)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if Path(LOGO_PATH).exists():
        hp.add_run().add_picture(LOGO_PATH, width=Inches(1.8))
    else:
        r = hp.add_run("Quarta Acompañamiento Legal S.A.S.")
        r.font.name = "Calibri"; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A); r.font.bold = True

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run("Quarta – Acompañamiento Legal S.A.S.\nwww.quarta.co  |  Tel: +57 3002766132\nCl. 5 Sur #43C-80 piso 8 NEWO On Going  |  Medellín, Colombia")
    fr.font.name = "Calibri"; fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # Barra dorada
    tb = doc.add_table(rows=1, cols=1)
    tc = tb.cell(0, 0)._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), COLOR_GOLD)
    tcPr.append(shd)
    tb.rows[0].height = Cm(0.25)
    doc.add_paragraph()

    hoy = datetime.now()
    meses = ["enero","febrero","marzo","abril","mayo","junio",
             "julio","agosto","septiembre","octubre","noviembre","diciembre"]
    pf = doc.add_paragraph(f"Medellín, {hoy.day} de {meses[hoy.month-1]} de {hoy.year}")
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.runs[0].font.name = "Cambria"; pf.runs[0].font.size = Pt(11)
    doc.add_paragraph()

    for linea in respuesta_agente.strip().split("\n"):
        linea = linea.rstrip()
        if not linea:
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        es_titulo = linea.isupper() and len(linea) > 4
        run = p.add_run(linea)
        run.font.name = "Cambria"; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        if es_titulo:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(14)

    for _ in range(3):
        doc.add_paragraph()

    pf2 = doc.add_paragraph()
    pf2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rf = pf2.add_run("Andrés Felipe Cardona Mesa\nAbogado — Quarta Acompañamiento Legal S.A.S.\nT.P. 400.075 del C.S.J.")
    rf.font.name = "Cambria"; rf.font.size = Pt(11)
    rf.font.color.rgb = RGBColor(0x1B, 0x2A, 0x4A)

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = re.sub(r"[^\w\s-]", "", nombre_proceso).strip().replace(" ", "_")
    out = Path("./documentos_generados")
    out.mkdir(exist_ok=True)
    ruta = out / f"Escrito_{safe}_{ts}.docx"
    doc.save(str(ruta))
    return str(ruta)


def generar_documento(solicitud: str, respuesta_agente: str,
                      nombre_proceso: str = "Proceso") -> dict:
    tipo = detectar_tipo_documento(solicitud)
    if tipo == "excel":
        ruta = generar_excel_calificacion(solicitud, respuesta_agente, nombre_proceso)
        return {"tipo": "excel", "ruta": ruta, "nombre_archivo": Path(ruta).name}
    elif tipo == "word":
        ruta = generar_word_escrito(respuesta_agente, nombre_proceso)
        return {"tipo": "word", "ruta": ruta, "nombre_archivo": Path(ruta).name}
    return {"tipo": "texto"}
